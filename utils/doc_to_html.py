from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docxtpl import DocxTemplate
import re

def docx_to_html(file_path):
    """
    Convert DOCX to HTML while preserving formatting, original element order, headers and footers.
    """
    template_doc = DocxTemplate(file_path)
    doc = template_doc.get_docx()
    html_content = ""

    def process_run_formatting(run):
        """Process text formatting for a run"""
        text = run.text
        if text.strip():
            # Preserve placeholders
            placeholder_matches = re.findall(r"{{(.*?)}}", text)
            if placeholder_matches:
                return text
            
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            
            if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                rgb = run.font.color.rgb
                text = f'<span style="color: rgb({rgb.r},{rgb.g},{rgb.b})">{text}</span>'
            
            if run.font.size:
                size = run.font.size.pt
                text = f'<span style="font-size: {size}pt">{text}</span>'
        return text

    def process_paragraph(paragraph):
        """Process a paragraph with its formatting"""
        style = paragraph.style.name
        alignment = paragraph.alignment
        
        style_attr = ''
        if alignment is not None:
            align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
            if alignment in align_map:
                style_attr = f' style="text-align: {align_map[alignment]}"'
        
        tag = 'p'
        if 'Heading' in style:
            level = style[-1] if style[-1].isdigit() else '1'
            tag = f'h{level}'
        
        content = ''.join(process_run_formatting(run) for run in paragraph.runs)
        return f"<{tag}{style_attr}>{content}</{tag}>"

    def process_table(table):
        """Process a table with its formatting"""
        table_html = '<table border="1" style="width:100%; border-collapse: collapse; margin: 10px 0;">'
        for row in table.rows:
            table_html += "<tr>"
            for cell in row.cells:
                table_html += '<td style="border: 1px solid #ddd; padding: 8px;">'
                for paragraph in cell.paragraphs:
                    table_html += process_paragraph(paragraph)
                table_html += "</td>"
            table_html += "</tr>"
        table_html += "</table>"
        return table_html

    def process_section(section):
        """Process headers and footers in a section"""
        section_html = ""
        
        # Process header
        if section.header:
            section_html += '<div class="header" style="border-bottom: 1px solid #ddd; margin-bottom: 20px; padding-bottom: 10px;">'
            # Process header paragraphs
            for paragraph in section.header.paragraphs:
                section_html += process_paragraph(paragraph)
            # Process header tables
            for table in section.header.tables:
                section_html += process_table(table)
            section_html += '</div>'
        
        # Process footer
        if section.footer:
            section_html += '<div class="footer" style="border-top: 1px solid #ddd; margin-top: 20px; padding-top: 10px;">'
            # Process footer paragraphs
            for paragraph in section.footer.paragraphs:
                section_html += process_paragraph(paragraph)
            # Process footer tables
            for table in section.footer.tables:
                section_html += process_table(table)
            section_html += '</div>'
            
        return section_html

    # Process each section (for headers and footers)
    for section in doc.sections:
        html_content += process_section(section)

    # Process main document content
    for element in doc.element.body:
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, doc)
            html_content += process_paragraph(paragraph)
        elif isinstance(element, CT_Tbl):
            table = Table(element, doc)
            html_content += process_table(table)

    return html_content